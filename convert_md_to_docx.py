#!/usr/bin/env python3
"""
Convert Markdown to DOCX format with image and hyperlink support.
"""

import re
import urllib.parse
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add blue color and underline
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def process_inline_formatting(paragraph, text, doc_dir):
    """Process inline markdown formatting like bold, italic, code, links."""
    # Pattern for links: [text](url)
    link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'

    # Split text by links and process
    parts = []
    last_end = 0

    for match in re.finditer(link_pattern, text):
        if match.start() > last_end:
            parts.append(('text', text[last_end:match.start()]))
        parts.append(('link', match.group(1), match.group(2)))
        last_end = match.end()

    if last_end < len(text):
        parts.append(('text', text[last_end:]))

    for part in parts:
        if part[0] == 'text':
            content = part[1]
            # Process bold - find **text** patterns
            bold_parts = re.split(r'\*\*([^*]+)\*\*', content)
            for i, bp in enumerate(bold_parts):
                if i % 2 == 1:  # Bold text
                    run = paragraph.add_run(bp)
                    run.bold = True
                elif bp:  # Regular text
                    # Process inline code
                    code_parts = re.split(r'`([^`]+)`', bp)
                    for j, cp in enumerate(code_parts):
                        if j % 2 == 1:  # Code text
                            run = paragraph.add_run(cp)
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                        elif cp:
                            paragraph.add_run(cp)
        elif part[0] == 'link':
            link_text = part[1]
            link_url = part[2]
            add_hyperlink(paragraph, link_text, link_url)


def convert_md_to_docx(md_path, docx_path):
    """Convert a markdown file to DOCX format."""

    md_path = Path(md_path)
    doc_dir = md_path.parent

    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Process content line by line
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_content = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                in_code_block = False
                if code_block_content:
                    para = doc.add_paragraph()
                    para.style = 'No Spacing'
                    for code_line in code_block_content:
                        run = para.add_run(code_line + '\n')
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    code_block_content = []
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # Handle images: ![alt](path)
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', line)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            # Decode URL encoding (e.g., %20 -> space)
            img_path = urllib.parse.unquote(img_path)

            # Resolve relative path
            full_img_path = doc_dir / img_path

            if full_img_path.exists():
                # Add image
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(str(full_img_path), width=Inches(6.0))

                # Add caption if alt text exists
                if alt_text:
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = caption.add_run(f"Figure: {alt_text}")
                    run.italic = True
                    run.font.size = Pt(10)
            else:
                # Image not found, add placeholder text
                para = doc.add_paragraph()
                para.add_run(f"[Image not found: {img_path}]")

            i += 1
            continue

        # Handle tables
        if '|' in line and not line.startswith('```'):
            stripped = line.strip()
            if stripped.startswith('|') and stripped.endswith('|'):
                if not in_table:
                    in_table = True
                    table_rows = []

                # Skip separator rows (|---|---|)
                if re.match(r'^\|[\s\-:|]+\|$', stripped):
                    i += 1
                    continue

                # Parse table row
                cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
                table_rows.append(cells)
                i += 1
                continue

        # End table if we were in one
        if in_table and ('|' not in line or line.startswith('```')):
            in_table = False
            if table_rows:
                num_cols = max(len(row) for row in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Table Grid'

                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            # Clean up markdown formatting
                            cell_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_data)
                            cell_text = re.sub(r'`([^`]+)`', r'\1', cell_text)
                            cell.text = cell_text

                doc.add_paragraph()
                table_rows = []

        # Handle horizontal rules
        if line.strip() == '---' or line.strip() == '***':
            doc.add_paragraph('_' * 50)
            i += 1
            continue

        # Handle headings
        if line.startswith('#'):
            match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if match:
                level = len(match.group(1))
                text = match.group(2)

                if level == 1:
                    para = doc.add_heading(text, level=0)
                else:
                    para = doc.add_heading(text, level=min(level, 9))

                i += 1
                continue

        # Handle bullet lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            para = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(para, text, doc_dir)
            i += 1
            continue

        # Handle numbered lists
        match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if match:
            text = match.group(2)
            para = doc.add_paragraph(style='List Number')
            process_inline_formatting(para, text, doc_dir)
            i += 1
            continue

        # Handle italic text (lines starting with *)
        if line.strip().startswith('*') and not line.strip().startswith('**') and line.strip().endswith('*'):
            text = line.strip()[1:-1]
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.italic = True
            i += 1
            continue

        # Handle regular paragraphs
        if line.strip():
            para = doc.add_paragraph()
            process_inline_formatting(para, line, doc_dir)

        i += 1

    # Handle any remaining table
    if in_table and table_rows:
        num_cols = max(len(row) for row in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        table.style = 'Table Grid'

        for row_idx, row_data in enumerate(table_rows):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    cell_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_data)
                    cell_text = re.sub(r'`([^`]+)`', r'\1', cell_text)
                    cell.text = cell_text

    # Save document
    doc.save(docx_path)
    print(f"Document saved to: {docx_path}")


if __name__ == '__main__':
    md_path = Path('/home/user/lowcost-star-tracker/docs/LowCost_StarTracker_Technical_Paper.md')
    docx_path = Path('/home/user/lowcost-star-tracker/docs/LowCost_StarTracker_Technical_Paper.docx')

    convert_md_to_docx(md_path, docx_path)
